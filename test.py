import docx
from docx.shared import Pt
#creating a New Document
doc = docx.Document()

# Adding Headings
doc.add_heading('Heading for the Document',0)
doc.add_heading('1st Level',1)
doc.add_heading('2nd Level',2)
doc.add_heading('3rd level',3)
doc.add_heading('4 Th level',4)
doc.add_heading('5 Th level',5)
doc.add_heading('6 Th level',6)
doc.add_heading('7 Th level',7)
doc.add_heading('8 Th level',8)
doc.add_heading('9 Th level',9)

# Creating Paragraph
dp = doc.add_paragraph('Paragraph 1')
#Adding Styles to text in Paragraphs
dp.add_run('hey there B').bold = True
dp.add_run('  and  ')
dp.add_run('These are italic ').italic = True
dp.add_run('These are U ').underline = True
dp.add_run('all caps').font.all_caps=True
dp.add_run('Bold').font.bold=True
dp.add_run('double strike').font.double_strike=True
dp.add_run('Emboss').font.emboss=True
dp.add_run('imprint').font.imprint=True
dp.add_run('Italic').font.italic=True
dp.add_run('Math').font.math=True
dp.add_run('all caps').font.name='Comic Sans MS'
dp.add_run('Outline').font.outline=True
dp.add_run('Shadow').font.shadow=True
dp.add_run('Size 23').font.size=Pt(23)
dp.add_run('small caqPS').font.small_caps=True
dp.add_run('strike').font.strike=True
dp.add_run('sub').font.subscript=True
dp.add_run('sup').font.superscript=True




#Going to new page
doc.add_page_break()
doc.add_paragraph("Hey 2nd Paragraph")
doc.add_heading('1',1)
#Saving the Document
doc.save('C:\\Users\\new\\Desktop\\mc1.docx')